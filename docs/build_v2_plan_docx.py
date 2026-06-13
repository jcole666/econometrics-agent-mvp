from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "计量建模Agent_MVP_V2开发分工与推进方案.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_east_asian_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_run(paragraph, text: str, bold: bool = False, color=None, size=None):
    run = paragraph.add_run(text)
    set_east_asian_font(run)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 22, DARK_BLUE, 0, 8),
        ("Subtitle", 11, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, item)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        add_run(p, item)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for line in text.splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(51, 51, 51)
    p_format = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_format.append(shd)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, label + "：", bold=True, color=DARK_BLUE)
    add_run(p, text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hdr[i].paragraphs[0], header, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            if i == 0:
                add_run(p, value, bold=True, color=DARK_BLUE)
            else:
                add_run(p, value)
    set_table_geometry(table, widths_dxa)


def build_doc() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "计量建模 Agent MVP V2", bold=True, color=DARK_BLUE, size=22)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "开发分工与推进方案")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, "适用对象：6 人小组 / 中期后续开发 / Agent 架构讨论", color=MUTED)

    add_callout(
        doc,
        "一句话目标",
        "在现有 MVP 基础上，从“模型推荐 + 代码生成工具”升级为“能对话、能辅助变量识别、能运行基础模型、能解释结果的计量建模 Agent 原型”。",
    )

    add_heading(doc, "1. V2 阶段目标", 1)
    p = doc.add_paragraph()
    add_run(
        p,
        "当前项目已经完成轻量级 MVP：上传 CSV / Excel 数据、识别字段信息、根据研究问题和变量配置推荐计量模型、生成 Python 计量建模代码模板，并可选调用华为云 MaaS 校对推荐结果。",
    )
    add_bullets(
        doc,
        [
            "加入大模型对话区，让用户能围绕数据、模型和代码继续追问。",
            "自动识别 Y / X 变量，降低用户手动填写门槛。",
            "保留用户确认或修改变量的入口，避免计量建模变成黑箱。",
            "至少实现 OLS 模型真实运行，返回基础回归结果。",
            "使用大模型解释推荐理由、代码和回归结果。",
        ],
    )

    add_heading(doc, "2. 推荐开发路线", 1)
    add_heading(doc, "第一阶段：大模型对话区", 2)
    add_bullets(
        doc,
        [
            "在现有页面中增加“和大模型对话”的区域。",
            "支持用户追问为什么推荐某个模型、变量如何理解、代码每行含义、模型风险和后续分析建议。",
            "第一阶段重点是让系统更像 Agent，但暂时不急着做复杂任务规划。",
        ],
    )

    add_heading(doc, "第二阶段：自动识别变量", 2)
    add_code_block(
        doc,
        "用户上传数据\n用户输入研究问题\nAgent 自动识别 Y / X / 控制变量\n用户确认或修改变量\n系统推荐模型",
    )
    p = doc.add_paragraph()
    add_run(p, "注意：不建议完全删除 Y / X 输入框。更合理的方式是 Agent 自动填入，用户仍然可以修改。计量建模中变量选择很敏感，保留人工确认更稳妥。")

    add_heading(doc, "第三阶段：模型真实运行", 2)
    add_code_block(
        doc,
        "用户上传数据\n选择或自动识别 Y / X\n点击运行模型\n后端使用 statsmodels 执行 OLS\n返回系数、p 值、R²、样本量等结果\n大模型解释回归结果",
    )

    add_heading(doc, "第四阶段：完整 Agent 化", 2)
    add_bullets(
        doc,
        [
            "多轮对话记忆。",
            "自动拆解任务。",
            "自动执行代码。",
            "根据报错反思和修复。",
            "用户修改变量后重新规划。",
            "生成 Markdown / Word 分析报告。",
            "保存每次分析过程。",
        ],
    )

    add_heading(doc, "3. 六人分工方案", 1)
    add_table(
        doc,
        ["角色", "主要职责", "入手位置", "建议交付物"],
        [
            [
                "前端交互负责人",
                "增加大模型对话区、自动识别变量按钮、运行模型按钮、回归结果和模型解释展示区域。",
                "先看 app/main.py 中的 CHINESE_DEMO_HTML；不会代码也可以先画页面草图。",
                "页面草图、修改后的演示页、对话区、变量识别展示区、结果展示区。",
            ],
            [
                "后端接口负责人",
                "新增 /chat、/infer-variables、/run-model 接口，设计请求和返回字段，打通前后端。",
                "参考 app/main.py 中的 /profile-data 和 /recommend-model。",
                "三个接口初版、接口测试记录、前后端联调记录。",
            ],
            [
                "大模型 Prompt 负责人",
                "编写对话助手、变量识别、模型解释、代码解释、结果解释 prompt。",
                "先整理 prompts.md，每个 prompt 配 2-3 个测试问题。",
                "Prompt 文档、测试样例、效果记录。",
            ],
            [
                "计量模型规则负责人",
                "梳理 OLS、Logit、DID、RDD、IV-2SLS、面板模型适用条件、变量要求和检查项。",
                "先看 app/services/model_selector.py；不会代码可先写中文规则表。",
                "模型选择规则文档、变量要求、检查清单、典型案例。",
            ],
            [
                "模型执行与结果解析负责人",
                "先实现 OLS 真实运行，返回系数、标准误、p 值、R²、样本量等结果。",
                "先看 app/services/code_generator.py；可以新增 model_runner.py。",
                "OLS 运行函数、/run-model 接口、回归结果 JSON。",
            ],
            [
                "测试样例与集成负责人",
                "准备测试数据和研究问题，记录推荐结果、运行结果和 Demo 流程。",
                "先用 Excel 或 CSV 做 20-50 行小数据，每个数据配一个研究问题。",
                "测试数据集、测试问题清单、测试记录表、最终 Demo 脚本。",
            ],
        ],
        [1500, 3300, 2300, 2260],
    )

    add_heading(doc, "4. 各角色详细任务", 1)
    role_details = [
        (
            "前端交互负责人",
            [
                "在现有页面中增加“大模型对话区”。",
                "增加“自动识别变量”按钮，并把结果填入 Y / X 输入框。",
                "增加“运行模型”按钮。",
                "增加“回归结果”和“模型解释”展示区域。",
                "优化页面布局，让演示流程更清楚。",
            ],
        ),
        (
            "后端接口负责人",
            [
                "新增 /chat 接口，用于用户和大模型对话。",
                "新增 /infer-variables 接口，用于自动识别 Y / X。",
                "新增 /run-model 接口，用于运行模型。",
                "设计每个接口的入参和返回字段。",
                "将前端页面与后端接口打通。",
            ],
        ),
        (
            "大模型 Prompt 负责人",
            [
                "编写大模型对话区 system prompt。",
                "编写变量识别 prompt。",
                "编写模型推荐解释 prompt。",
                "编写代码解释 prompt。",
                "编写回归结果解释 prompt。",
                "测试不同问法下回答是否稳定。",
            ],
        ),
        (
            "计量模型规则负责人",
            [
                "整理 OLS、Logit、DID、RDD、IV-2SLS、面板模型的适用条件。",
                "整理每个模型需要哪些变量。",
                "补充每个模型的建模前检查项。",
                "检查现有规则是否合理。",
                "给模型选择模块提供中文规则说明。",
            ],
        ),
        (
            "模型执行与结果解析负责人",
            [
                "先实现 OLS 的真实运行。",
                "后端读取用户上传的数据。",
                "根据 Y / X 调用 statsmodels 执行 OLS。",
                "返回系数、标准误、p 值、R²、样本量等结果。",
                "后续再支持 Logit、DID、RDD、IV-2SLS。",
            ],
        ),
        (
            "测试样例与集成负责人",
            [
                "准备 4-6 个小型 CSV 数据集。",
                "每个数据集配一个研究问题。",
                "测试系统是否推荐正确模型。",
                "测试生成代码是否合理。",
                "测试 OLS 是否能真实运行。",
                "整理最终 Demo 流程。",
            ],
        ),
    ]
    for role, tasks in role_details:
        add_heading(doc, role, 2)
        add_bullets(doc, tasks)

    add_heading(doc, "5. 推荐时间安排", 1)
    add_table(
        doc,
        ["阶段", "重点工作"],
        [
            [
                "第一周：对话和规则基础",
                "前端加入大模型对话区；后端实现 /chat；Prompt 负责人写对话助手 prompt；计量规则负责人整理模型规则；执行负责人研究 OLS；测试负责人准备 OLS 和 DID 样例。",
            ],
            [
                "第二周：变量识别和 OLS 运行",
                "实现 /infer-variables；自动识别出的 Y / X 填入页面；接入大模型解释推荐理由；实现 OLS 真实运行；返回基础回归结果；准备完整 Demo 流程。",
            ],
            [
                "第三周：扩展模型和优化演示",
                "尝试扩展 Logit 或 DID；优化页面展示；补充测试样例；整理架构图；整理最终说明文档；准备汇报话术。",
            ],
        ],
        [2300, 7060],
    )

    add_heading(doc, "6. 对没有代码基础成员的入手建议", 1)
    add_heading(doc, "第一步：先理解当前页面流程", 2)
    add_code_block(doc, "01 上传数据并识别字段\n02 输入研究问题并推荐模型\n03 支持的模型类型")
    add_bullets(
        doc,
        [
            "能说清楚用户输入了什么。",
            "能说清楚系统识别了什么。",
            "能说清楚系统推荐了什么模型。",
            "能说清楚系统生成了什么代码。",
        ],
    )
    add_heading(doc, "第二步：先写中文规则和样例", 2)
    add_bullets(
        doc,
        [
            "写模型适用条件。",
            "写研究问题样例。",
            "写 Prompt。",
            "做测试 CSV。",
            "记录测试结果。",
        ],
    )
    add_heading(doc, "第三步：跟着现有代码改小地方", 2)
    add_bullets(
        doc,
        [
            "改页面文字。",
            "改按钮名称。",
            "改模型说明。",
            "补充检查项。",
            "添加新的研究问题样例。",
            "在规则文档中增加关键词。",
        ],
    )
    add_heading(doc, "第四步：逐渐参与接口和功能实现", 2)
    add_bullets(
        doc,
        [
            "新增接口。",
            "调 prompt。",
            "调模型运行函数。",
            "调前端按钮。",
            "整理接口测试。",
        ],
    )

    add_heading(doc, "7. 会上推荐说法", 1)
    add_callout(
        doc,
        "建议表达",
        "我们后续建议按模块分工，不只是做 PPT 或写材料。6 个人分别负责前端交互、后端接口、大模型 Prompt、计量模型规则、模型执行、测试样例与集成。第一阶段先加大模型对话区和变量自动识别，让系统更像一个 Agent；第二阶段实现 OLS 的真实运行和结果解释；第三阶段再逐步扩展 DID、RDD、IV 等复杂模型。这样每个人都有明确产出，也能保证项目一步一步可演示。",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(footer, "计量建模 Agent MVP V2 开发分工与推进方案", color=MUTED, size=9)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
